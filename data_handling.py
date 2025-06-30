import streamlit as st
import json
from docx import Document
import fitz  # PyMuPDF
import os
from utils import generate_chat_title
import markdown
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datatypes.CV_obj import CV
try:
    master_cv_list = [f for f in os.listdir("MasterCV") if (f.endswith(".pdf") or f.endswith(".docx"))]
except:
    master_cv_list = []

save_dir = "masterCV"
os.makedirs(save_dir, exist_ok=True)

def read_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text


def read_docx(file_path):
    doc = Document(file_path)
    return "\\n".join(para.text for para in doc.paragraphs)

def save_chat_history(history, model):
    title = generate_chat_title(history, model)
    filename = f"chat history/{title}.json"
    chat_data = {
        "title": title,
        "history": history
    }
    with open(filename, 'w') as f:
        json.dump(chat_data, f)

def load_chat_history(filename):
    try:
        with open(filename, 'r') as f:
            chat_data = json.load(f)
        return chat_data["history"]
    except FileNotFoundError:
        return []

def create_docx(text):
    # Create a new document
    document = Document()
    text = text.replace("```", "").replace("markdown","")
    document.add_paragraph(text)
    if st.session_state['company_name'] is not None:
        filename = st.session_state['user_name'] + "_" + st.session_state['company_name']
    else:
        filename = st.session_state['user_name'] + "edited"
    document.save(filename)
    return filename

def read_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return read_docx(file_path)
    elif ext == ".pdf":
        return read_pdf(file_path)
    else:
        raise ValueError("Unsupported file type")

def create_docx_from_CV(cv_data: CV) -> str:
    document = Document()

    # Full Name
    fullname_paragraph = document.add_paragraph()
    fullname_run = fullname_paragraph.add_run(cv_data.full_name)
    fullname_run.bold = True
    fullname_run.font.size = Pt(12)
    fullname_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Details
    contact_paragraph = document.add_paragraph()
    contact_details = f"{cv_data.linkedin} {cv_data.contact_details.email_id} {cv_data.contact_details.phone_number if cv_data.contact_details.phone_number is not None else ''}"
    if cv_data.github:
        contact_details += f" | {cv_data.github}"
    contact_run = contact_paragraph.add_run(contact_details)
    contact_run.font.size = Pt(10)
    contact_run.underline = True
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Personal Profile
    if cv_data.personal_profile:
        document.add_paragraph().add_run("Personal Profile").bold = True
        document.add_paragraph(cv_data.personal_profile).add_run().font.size = Pt(10)

    # Skills
    if cv_data.skills:
        document.add_paragraph().add_run("Skills").bold = True
        for skill_cat in cv_data.skills:
            skill_cat_paragraph = document.add_paragraph()
            skill_cat_run = skill_cat_paragraph.add_run(f"{skill_cat.skill_subcategory}: ")
            skill_cat_run.bold = True
            skill_cat_run.font.size = Pt(10)
            skills_list_run = skill_cat_paragraph.add_run(", ".join(skill_cat.skill_list))
            skills_list_run.font.size = Pt(10)

    # Experience
    if cv_data.experience:
        document.add_paragraph().add_run("Experience").bold = True
        for exp in cv_data.experience:
            exp_heading = document.add_paragraph()
            company_date_run = exp_heading.add_run(f"{exp.company_name} ({exp.date_from} - {exp.date_to})")
            company_date_run.bold = True
            company_date_run.font.size = Pt(11)
            exp_heading.add_run("\n") # Add a newline
            role_run = exp_heading.add_run(f"{exp.role} {exp.location if exp.location is not None else ''}")
            role_run.font.size = Pt(11) # Keep the same size, but not bold
            for responsibility in exp.responsibilities:
                document.add_paragraph(responsibility, style='List Bullet').add_run().font.size = Pt(10)

    # Education
    if cv_data.education:
        document.add_paragraph().add_run("Education").bold = True
        for edu in cv_data.education:
            edu_heading = document.add_paragraph()
            edu_heading_run = edu_heading.add_run(f"{edu.degree} from {edu.university} ({edu.date_from} - {edu.date_to})")
            edu_heading_run.bold = True
            edu_heading_run.font.size = Pt(11)
            if edu.grade:
                document.add_paragraph(f"Grade: {edu.grade}").add_run().font.size = Pt(10)

    # Projects
    if cv_data.projects:
        document.add_paragraph().add_run("Projects").bold = True
        for proj in cv_data.projects:
            proj_heading = document.add_paragraph()
            proj_heading_run = proj_heading.add_run(f"{proj.project_title}")
            proj_heading_run.bold = True
            proj_heading_run.font.size = Pt(11)
            if proj.description:
                for detail in proj.description:
                    document.add_paragraph(detail, style='List Bullet').add_run().font.size = Pt(10)

    # Save the document
    filename = f"{cv_data.full_name.replace(' ', '_')}_{st.session_state['company_name']}_CV.docx"
    document.save(filename)
    return filename
